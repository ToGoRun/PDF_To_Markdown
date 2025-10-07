"""
PDF to Markdown Converter with Preview
使用marker-pdf库实现PDF到Markdown的转换，带有逐页预览功能
"""

import sys
import os
from pathlib import Path

# 确保使用虚拟环境
script_dir = Path(__file__).parent
venv_python = script_dir / "venv" / "Scripts" / "python.exe"

if venv_python.exists() and sys.executable != str(venv_python):
    # 如果虚拟环境存在且当前不是在虚拟环境中运行，重新启动
    import subprocess
    subprocess.Popen([str(venv_python), __file__])
    sys.exit(0)

import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from tkinter.scrolledtext import ScrolledText
import threading
import fitz  # PyMuPDF用于PDF预览
from PIL import Image, ImageTk
import io
import shutil
import re

class PDFtoMarkdownConverter:
    def __init__(self, root):
        self.root = root
        self.root.title("PDF转Markdown转换器 (Marker-PDF)")
        self.root.geometry("1400x800")

        self.pdf_path = None
        self.markdown_content = ""
        self.pdf_document = None
        self.current_page = 0
        self.total_pages = 0
        self.page_markdown_map = {}  # 存储每页对应的markdown内容
        self.rendered_output = None  # 存储marker的完整输出
        self.extracted_images = {}  # 存储提取的图片 {block_id: base64_data}

        # 模型缓存目录
        self.models_dir = script_dir / "models"
        self.models_dir.mkdir(exist_ok=True)

        # 检测GPU设备
        self.device_info = self._detect_device()

        # 性能优化选项
        self.use_ocr = tk.BooleanVar(value=False)  # 默认不使用OCR (快速模式)
        self.high_quality_mode = tk.BooleanVar(value=False)  # 高质量模式

        self.setup_ui()

    def _detect_device(self):
        """检测可用的计算设备"""
        try:
            import torch

            device_info = {
                'device': 'cpu',
                'name': 'CPU',
                'cuda_available': False,
                'gpu_name': None
            }

            if torch.cuda.is_available():
                device_info['cuda_available'] = True
                device_info['device'] = 'cuda'
                device_info['gpu_name'] = torch.cuda.get_device_name(0)
                device_info['name'] = f"GPU: {device_info['gpu_name']}"

                # 检查是否是RTX系列显卡
                gpu_name_lower = device_info['gpu_name'].lower()
                if 'rtx' in gpu_name_lower or 'geforce' in gpu_name_lower or 'quadro' in gpu_name_lower:
                    print(f"检测到NVIDIA GPU: {device_info['gpu_name']}")
                    print(f"CUDA版本: {torch.version.cuda}")
                    print(f"将使用GPU加速处理")
                else:
                    print(f"检测到GPU: {device_info['gpu_name']}")
            else:
                print("未检测到可用GPU，将使用CPU处理（速度较慢）")

            return device_info

        except ImportError:
            print("PyTorch未安装，将使用CPU处理")
            return {
                'device': 'cpu',
                'name': 'CPU',
                'cuda_available': False,
                'gpu_name': None
            }
        except Exception as e:
            print(f"设备检测失败: {str(e)}，将使用CPU处理")
            return {
                'device': 'cpu',
                'name': 'CPU',
                'cuda_available': False,
                'gpu_name': None
            }

    def setup_ui(self):
        # 顶部工具栏
        toolbar = ttk.Frame(self.root)
        toolbar.pack(side=tk.TOP, fill=tk.X, padx=5, pady=5)

        ttk.Button(toolbar, text="选择PDF文件", command=self.select_pdf).pack(side=tk.LEFT, padx=5)
        ttk.Button(toolbar, text="开始转换", command=self.convert_pdf).pack(side=tk.LEFT, padx=5)
        ttk.Button(toolbar, text="导出Markdown", command=self.export_markdown).pack(side=tk.LEFT, padx=5)
        ttk.Button(toolbar, text="导出Word", command=self.export_word).pack(side=tk.LEFT, padx=5)

        # 性能选项
        ttk.Separator(toolbar, orient=tk.VERTICAL).pack(side=tk.LEFT, padx=10, fill=tk.Y)
        ttk.Checkbutton(toolbar, text="使用OCR (慢)", variable=self.use_ocr,
                       command=self._update_mode_tooltip).pack(side=tk.LEFT, padx=5)
        ttk.Checkbutton(toolbar, text="高质量模式", variable=self.high_quality_mode).pack(side=tk.LEFT, padx=5)

        self.status_label = ttk.Label(toolbar, text=f"请选择PDF文件 | 设备: {self.device_info['name']} | 模式: 快速")
        self.status_label.pack(side=tk.LEFT, padx=20)

        # 进度条
        self.progress = ttk.Progressbar(toolbar, mode='indeterminate', length=200)
        self.progress.pack(side=tk.RIGHT, padx=5)
        
        # 主内容区域 - 分为左右两栏
        main_frame = ttk.Frame(self.root)
        main_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        # 左侧 - PDF预览
        left_frame = ttk.LabelFrame(main_frame, text="PDF预览")
        left_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(0, 5))
        
        # PDF画布
        self.pdf_canvas = tk.Canvas(left_frame, bg='white')
        self.pdf_canvas.pack(fill=tk.BOTH, expand=True)
        
        # 右侧 - Markdown预览
        right_frame = ttk.LabelFrame(main_frame, text="Markdown预览")
        right_frame.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True)
        
        self.markdown_text = ScrolledText(right_frame, wrap=tk.WORD, font=('Consolas', 10))
        self.markdown_text.pack(fill=tk.BOTH, expand=True)
        
        # 底部导航栏
        nav_frame = ttk.Frame(self.root)
        nav_frame.pack(side=tk.BOTTOM, fill=tk.X, padx=5, pady=5)
        
        ttk.Button(nav_frame, text="◀ 上一页", command=self.prev_page).pack(side=tk.LEFT, padx=5)
        
        self.page_label = ttk.Label(nav_frame, text="0 / 0")
        self.page_label.pack(side=tk.LEFT, padx=20)
        
        ttk.Button(nav_frame, text="下一页 ▶", command=self.next_page).pack(side=tk.LEFT, padx=5)
        
        # 页面跳转
        ttk.Label(nav_frame, text="跳转到页面:").pack(side=tk.LEFT, padx=(50, 5))
        self.page_entry = ttk.Entry(nav_frame, width=10)
        self.page_entry.pack(side=tk.LEFT, padx=5)
        ttk.Button(nav_frame, text="跳转", command=self.jump_to_page).pack(side=tk.LEFT, padx=5)
        
    def _update_mode_tooltip(self):
        """更新模式提示"""
        mode = "OCR模式 (慢但适合扫描PDF)" if self.use_ocr.get() else "快速模式 (使用PDF文本)"
        current_text = self.status_label.cget("text")
        parts = current_text.split("|")
        if len(parts) >= 2:
            self.status_label.config(text=f"{parts[0]}| {parts[1]}| 模式: {mode}")

    def select_pdf(self):
        """选择PDF文件"""
        file_path = filedialog.askopenfilename(
            title="选择PDF文件",
            filetypes=[("PDF文件", "*.pdf"), ("所有文件", "*.*")]
        )

        if file_path:
            self.pdf_path = file_path
            mode = "OCR" if self.use_ocr.get() else "快速"
            self.status_label.config(text=f"已选择: {os.path.basename(file_path)} | 设备: {self.device_info['name']} | 模式: {mode}")
            self.load_pdf_preview()
            
    def load_pdf_preview(self):
        """加载PDF预览"""
        try:
            if self.pdf_document:
                self.pdf_document.close()
                
            self.pdf_document = fitz.open(self.pdf_path)
            self.total_pages = len(self.pdf_document)
            self.current_page = 0
            self.display_page(0)
            
        except Exception as e:
            messagebox.showerror("错误", f"加载PDF失败: {str(e)}")
            
    def display_page(self, page_num):
        """显示指定页面"""
        if not self.pdf_document or page_num >= self.total_pages:
            return
            
        try:
            page = self.pdf_document[page_num]
            
            # 获取页面图像
            zoom = 1.5  # 缩放比例
            mat = fitz.Matrix(zoom, zoom)
            pix = page.get_pixmap(matrix=mat)
            
            # 转换为PIL Image
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            
            # 调整大小以适应画布
            canvas_width = self.pdf_canvas.winfo_width()
            canvas_height = self.pdf_canvas.winfo_height()
            
            if canvas_width > 1 and canvas_height > 1:
                img.thumbnail((canvas_width - 20, canvas_height - 20), Image.Resampling.LANCZOS)
            
            # 显示图像
            photo = ImageTk.PhotoImage(img)
            self.pdf_canvas.delete("all")
            self.pdf_canvas.create_image(
                canvas_width // 2 if canvas_width > 1 else 200,
                canvas_height // 2 if canvas_height > 1 else 300,
                image=photo
            )
            self.pdf_canvas.image = photo  # 保持引用
            
            # 更新页面标签
            self.page_label.config(text=f"{page_num + 1} / {self.total_pages}")
            
            # 显示对应的Markdown内容
            if page_num in self.page_markdown_map:
                self.markdown_text.delete(1.0, tk.END)
                self.markdown_text.insert(1.0, self.page_markdown_map[page_num])
            
        except Exception as e:
            messagebox.showerror("错误", f"显示页面失败: {str(e)}")
            
    def prev_page(self):
        """上一页"""
        if self.current_page > 0:
            self.current_page -= 1
            self.display_page(self.current_page)
            
    def next_page(self):
        """下一页"""
        if self.current_page < self.total_pages - 1:
            self.current_page += 1
            self.display_page(self.current_page)
            
    def jump_to_page(self):
        """跳转到指定页面"""
        try:
            page_num = int(self.page_entry.get()) - 1
            if 0 <= page_num < self.total_pages:
                self.current_page = page_num
                self.display_page(self.current_page)
            else:
                messagebox.showwarning("警告", f"页面编号必须在 1 到 {self.total_pages} 之间")
        except ValueError:
            messagebox.showwarning("警告", "请输入有效的页面编号")
            
    def convert_pdf(self):
        """转换PDF到Markdown"""
        if not self.pdf_path:
            messagebox.showwarning("警告", "请先选择PDF文件")
            return
            
        # 在新线程中执行转换
        thread = threading.Thread(target=self._convert_thread)
        thread.daemon = True
        thread.start()
        
    def _convert_thread(self):
        """转换线程"""
        try:
            self.root.after(0, self.progress.start)
            self.root.after(0, self.status_label.config, {"text": "正在转换中..."})

            # 重要: 必须在导入marker之前设置环境变量
            # 设置模型缓存目录 - marker使用surya库,需要设置MODEL_CACHE_DIR
            os.environ['MODEL_CACHE_DIR'] = str(self.models_dir)

            # 设置设备 - 必须在导入前设置才能生效
            os.environ['TORCH_DEVICE'] = self.device_info['device']

            # 额外的缓存路径设置(可选,用于其他可能的模型)
            os.environ['TORCH_HOME'] = str(self.models_dir)
            os.environ['HF_HOME'] = str(self.models_dir / 'huggingface')
            os.environ['TRANSFORMERS_CACHE'] = str(self.models_dir / 'huggingface' / 'transformers')

            # 如果使用GPU,显示GPU信息
            if self.device_info['cuda_available']:
                try:
                    import torch
                    print(f"\n{'='*50}")
                    print(f"使用GPU加速: {torch.cuda.get_device_name(0)}")
                    print(f"CUDA版本: {torch.version.cuda}")
                    print(f"GPU显存: {torch.cuda.get_device_properties(0).total_memory / 1024**3:.2f} GB")
                    print(f"{'='*50}\n")
                except Exception as e:
                    print(f"获取GPU信息失败: {e}")

            # 导入marker库 - 必须在环境变量设置之后
            try:
                from marker.converters.pdf import PdfConverter
                from marker.config.parser import ConfigParser
                from marker.models import create_model_dict
            except ImportError:
                self.root.after(0, messagebox.showerror, "错误",
                    "未安装marker-pdf库。请运行:\npip install marker-pdf")
                return

            # 加载模型 - 使用本地缓存目录
            device_msg = f"加载模型中... (使用{self.device_info['name']})"
            self.root.after(0, self.status_label.config, {"text": device_msg})

            # 创建模型字典 - 会自动使用TORCH_DEVICE环境变量
            print(f"正在加载模型到设备: {self.device_info['device']}")
            print(f"模型缓存目录: {self.models_dir}")
            artifact_dict = create_model_dict()

            # 性能优化配置
            use_ocr = self.use_ocr.get()
            high_quality = self.high_quality_mode.get()

            print(f"\n{'='*50}")
            print(f"转换模式配置:")
            print(f"  OCR模式: {'启用' if use_ocr else '禁用 (快速模式)'}")
            print(f"  高质量模式: {'启用' if high_quality else '禁用'}")
            print(f"{'='*50}\n")

            # 根据用户选择配置性能参数
            if use_ocr:
                # OCR模式 - 适合扫描PDF，速度较慢
                if high_quality:
                    # 高质量OCR模式
                    config_dict = {
                        "force_ocr": True,              # 强制使用OCR
                        "recognition_batch_size": 64,    # 适中的批次大小
                        "detection_batch_size": 10,
                        "disable_multiprocessing": True,
                    }
                    print("使用高质量OCR模式 (速度慢，质量高)")
                else:
                    # 快速OCR模式
                    config_dict = {
                        "force_ocr": True,              # 强制使用OCR
                        "recognition_batch_size": 128,   # 大批次提速
                        "detection_batch_size": 16,
                        "ocr_error_batch_size": 28,
                        "disable_ocr_math": True,       # 禁用数学符号识别加速
                        "disable_multiprocessing": True,
                    }
                    print("使用快速OCR模式 (速度中等，适合扫描PDF)")
            else:
                # 快速模式 - 使用PDF文本，不用OCR (推荐)
                config_dict = {
                    "disable_ocr": True,            # 禁用OCR，只用PDF文本
                    "pdftext_workers": 8,           # 并行提取文本
                    "disable_multiprocessing": True,
                }
                print("使用快速模式 (仅提取PDF文本，速度最快)")

            config = ConfigParser(config_dict)
            converter = PdfConverter(
                artifact_dict=artifact_dict,
                config=config.generate_config_dict()
            )

            # 执行转换
            self.root.after(0, self.status_label.config, {"text": "转换PDF中..."})
            rendered = converter(self.pdf_path)

            # 从rendered中提取markdown和图片
            from marker.output import text_from_rendered
            text, _, images = text_from_rendered(rendered)

            self.markdown_content = text
            self.rendered_output = rendered  # 保存完整输出
            self.extracted_images = images  # 保存提取的图片字典 {block_id: base64_data}

            # 调试信息
            print(f"\n提取的图片数量: {len(images) if images else 0}")
            if images:
                print(f"图片数据类型示例: {type(list(images.values())[0]) if images else 'N/A'}")
                print(f"图片ID列表: {list(images.keys())[:5]}...")  # 显示前5个

            # 简单分页处理 - 按页面分隔符分割
            # 注意: marker可能不会完美地按页面分割，这里做简单处理
            self._split_markdown_by_pages()
            
            self.root.after(0, self.progress.stop)
            self.root.after(0, self.status_label.config, {"text": "转换完成!"})
            self.root.after(0, messagebox.showinfo, "成功", "PDF转换完成!")
            
            # 显示第一页
            self.root.after(0, self.display_page, 0)
            
        except Exception as e:
            self.root.after(0, self.progress.stop)
            self.root.after(0, self.status_label.config, {"text": "转换失败"})
            self.root.after(0, messagebox.showerror, "错误", f"转换失败: {str(e)}")
            
    def _split_markdown_by_pages(self):
        """将Markdown内容按页面分割"""
        # 这是一个简化的实现
        # marker-pdf可能会在输出中包含页面标记或者我们需要按内容长度估算
        
        # 简单策略：平均分配内容到各页
        if self.total_pages > 0:
            lines = self.markdown_content.split('\n')
            lines_per_page = max(1, len(lines) // self.total_pages)
            
            for i in range(self.total_pages):
                start = i * lines_per_page
                end = start + lines_per_page if i < self.total_pages - 1 else len(lines)
                page_content = '\n'.join(lines[start:end])
                self.page_markdown_map[i] = page_content
        else:
            self.page_markdown_map[0] = self.markdown_content
            
    def _save_extracted_images(self, images_dir):
        """保存从marker提取的图片"""
        if not self.extracted_images:
            return 0

        images_dir = Path(images_dir)
        images_dir.mkdir(exist_ok=True)

        import base64
        image_count = 0
        self.image_filename_map = {}  # 记录block_id到实际文件名的映射

        # extracted_images 可能是不同的格式
        for block_id, image_data in self.extracted_images.items():
            try:
                image_bytes = None

                # 处理不同类型的图片数据
                if isinstance(image_data, str):
                    # 如果是字符串，尝试作为base64解码
                    try:
                        image_bytes = base64.b64decode(image_data)
                    except Exception:
                        print(f"跳过无效的base64数据 (block_id: {block_id})")
                        continue
                elif isinstance(image_data, bytes):
                    # 如果已经是字节数据，直接使用
                    image_bytes = image_data
                elif hasattr(image_data, 'save'):
                    # 如果是PIL Image对象
                    image_filename = f"{block_id}.png"
                    image_path = images_dir / image_filename
                    image_data.save(image_path)
                    self.image_filename_map[block_id] = image_filename
                    image_count += 1
                    continue
                else:
                    print(f"未知的图片数据类型: {type(image_data)} (block_id: {block_id})")
                    continue

                if image_bytes is None:
                    continue

                # 判断图片格式
                image_ext = "png"  # 默认使用png
                if len(image_bytes) >= 4:
                    if image_bytes[:4] == b'\xff\xd8\xff\xe0' or image_bytes[:4] == b'\xff\xd8\xff\xe1':
                        image_ext = "jpg"
                    elif len(image_bytes) >= 8 and image_bytes[:8] == b'\x89PNG\r\n\x1a\n':
                        image_ext = "png"
                    elif image_bytes[:2] == b'BM':
                        image_ext = "bmp"

                # 使用block_id作为文件名（保持原始ID，方便匹配）
                # marker生成的block_id通常是 _page_X_Picture_Y 或类似格式
                image_filename = f"{block_id}.{image_ext}"
                image_path = images_dir / image_filename

                with open(image_path, "wb") as img_file:
                    img_file.write(image_bytes)

                self.image_filename_map[block_id] = image_filename
                image_count += 1

                print(f"保存图片: {block_id} -> {image_filename}")

            except Exception as e:
                print(f"保存图片失败 (block_id: {block_id}): {str(e)}")
                import traceback
                traceback.print_exc()
                continue

        print(f"\n图片文件名映射表 (前5个):")
        for block_id, filename in list(self.image_filename_map.items())[:5]:
            print(f"  {block_id} -> {filename}")

        return image_count

    def _update_image_paths(self, markdown_content, images_folder_path):
        """更新markdown中的图片路径

        Args:
            images_folder_path: 图片文件夹路径（可以是相对路径或绝对路径）
        """
        import base64

        updated_content = markdown_content
        images_folder_path = Path(images_folder_path)

        print(f"\n更新图片路径:")
        print(f"图片文件夹路径: {images_folder_path}")
        print(f"是否为绝对路径: {images_folder_path.is_absolute()}")

        # 遍历所有提取的图片,替换markdown中的引用
        for block_id in self.extracted_images.keys():
            try:
                # 使用保存时记录的实际文件名
                if hasattr(self, 'image_filename_map') and block_id in self.image_filename_map:
                    image_filename = self.image_filename_map[block_id]
                else:
                    # 如果映射不存在，尝试检测格式
                    image_data = self.extracted_images[block_id]
                    image_ext = "png"  # 默认扩展名

                    try:
                        if isinstance(image_data, str):
                            image_bytes = base64.b64decode(image_data)
                        elif isinstance(image_data, bytes):
                            image_bytes = image_data
                        else:
                            image_bytes = None

                        if image_bytes and len(image_bytes) >= 4:
                            if image_bytes[:4] == b'\xff\xd8\xff\xe0' or image_bytes[:4] == b'\xff\xd8\xff\xe1':
                                image_ext = "jpg"
                            elif len(image_bytes) >= 8 and image_bytes[:8] == b'\x89PNG\r\n\x1a\n':
                                image_ext = "png"
                    except Exception:
                        pass

                    image_filename = f"{block_id}.{image_ext}"

                # 构建图片完整路径
                if images_folder_path.is_absolute():
                    # 绝对路径：完整路径
                    image_path = images_folder_path / image_filename
                    # 转换为URL格式的路径（Windows路径需要转换）
                    image_path_str = str(image_path).replace('\\', '/')
                else:
                    # 相对路径
                    image_path_str = f"{images_folder_path.name}/{image_filename}"

                # 查找并替换多种可能的图片引用格式
                # 先尝试精确匹配block_id
                pattern1 = rf'!\[([^\]]*)\]\({re.escape(block_id)}(?:\.[a-zA-Z]+)?\)'
                updated_content = re.sub(pattern1, rf'![\1]({image_path_str})', updated_content)

                # 再尝试匹配包含路径的引用
                pattern2 = rf'!\[([^\]]*)\]\([^\)]*{re.escape(block_id)}(?:\.[a-zA-Z]+)?\)'
                updated_content = re.sub(pattern2, rf'![\1]({image_path_str})', updated_content)

            except Exception as e:
                print(f"更新图片路径失败 (block_id: {block_id}): {str(e)}")
                continue

        return updated_content

    def export_markdown(self):
        """导出Markdown文件并提取图片"""
        if not self.markdown_content:
            messagebox.showwarning("警告", "请先转换PDF文件")
            return

        # 选择保存位置
        save_path = filedialog.asksaveasfilename(
            title="保存Markdown文件",
            defaultextension=".md",
            filetypes=[("Markdown文件", "*.md"), ("所有文件", "*.*")],
            initialfile=Path(self.pdf_path).stem + ".md" if self.pdf_path else "output.md"
        )

        if save_path:
            try:
                save_path = Path(save_path)
                images_dir = save_path.parent / f"{save_path.stem}_images"

                # 保存从marker提取的图片
                image_count = self._save_extracted_images(images_dir)

                # 更新markdown中的图片路径
                markdown_with_images = self._update_image_paths(self.markdown_content, images_dir.name)

                # 保存Markdown文件
                with open(save_path, 'w', encoding='utf-8') as f:
                    f.write(markdown_with_images)

                msg = f"Markdown文件已保存到:\n{save_path}"
                if image_count > 0:
                    msg += f"\n\n提取了 {image_count} 张图片到:\n{images_dir}"
                else:
                    msg += "\n\n注意: 未检测到图片或图片提取失败"
                messagebox.showinfo("成功", msg)
            except Exception as e:
                messagebox.showerror("错误", f"保存文件失败: {str(e)}")

    def export_word(self):
        """导出Word文件"""
        if not self.markdown_content:
            messagebox.showwarning("警告", "请先转换PDF文件")
            return

        # 选择保存位置
        save_path = filedialog.asksaveasfilename(
            title="保存Word文件",
            defaultextension=".docx",
            filetypes=[("Word文档", "*.docx"), ("所有文件", "*.*")],
            initialfile=Path(self.pdf_path).stem + ".docx" if self.pdf_path else "output.docx"
        )

        if save_path:
            try:
                save_path = Path(save_path)
                images_dir = save_path.parent / f"{save_path.stem}_images"

                print(f"\n{'='*60}")
                print(f"导出Word文档")
                print(f"{'='*60}")
                print(f"输出路径: {save_path}")
                print(f"图片目录: {images_dir}")

                # 保存图片
                print(f"\n步骤1: 保存图片...")
                print(f"提取的图片数量: {len(self.extracted_images)}")
                image_count = self._save_extracted_images(images_dir)
                print(f"成功保存图片: {image_count}")

                # 更新markdown中的图片路径
                print(f"\n步骤2: 更新Markdown中的图片路径...")
                # 对于Pandoc，使用绝对路径避免中文路径问题
                markdown_with_images = self._update_image_paths(self.markdown_content, str(images_dir))

                # 检查markdown中的图片引用
                import re
                img_refs = re.findall(r'!\[([^\]]*)\]\(([^\)]+)\)', markdown_with_images)
                print(f"Markdown中的图片引用数量: {len(img_refs)}")
                if img_refs:
                    print("图片引用示例:")
                    for alt, path in img_refs[:3]:  # 显示前3个
                        print(f"  - ![{alt}]({path})")

                # 转换Markdown到Word
                print(f"\n步骤3: 转换为Word文档...")
                self._convert_markdown_to_docx(markdown_with_images, save_path, images_dir)

                print(f"\n{'='*60}")

                msg = f"Word文件已保存到:\n{save_path}"
                if image_count > 0:
                    msg += f"\n\n提取了 {image_count} 张图片到:\n{images_dir}"
                else:
                    msg += "\n\n注意: 未检测到图片或图片提取失败"
                messagebox.showinfo("成功", msg)
            except Exception as e:
                import traceback
                traceback.print_exc()
                messagebox.showerror("错误", f"导出Word失败: {str(e)}\n\n提示: 可能需要安装Pandoc，详见README.md")

    def _convert_markdown_to_docx(self, markdown_content, output_path, images_dir=None):
        """将Markdown转换为Word文档"""
        try:
            # 方法1: 尝试使用pypandoc
            import pypandoc

            # 创建临时markdown文件
            temp_md = output_path.parent / f"temp_{output_path.stem}.md"
            with open(temp_md, 'w', encoding='utf-8') as f:
                f.write(markdown_content)

            print(f"临时Markdown文件: {temp_md}")
            print(f"图片目录: {images_dir}")

            # 使用pypandoc转换，指定资源路径
            extra_args = ['--standalone']
            if images_dir:
                # 添加资源路径，让Pandoc能找到图片
                extra_args.append(f'--resource-path={images_dir.parent}')

            pypandoc.convert_file(
                str(temp_md),
                'docx',
                outputfile=str(output_path),
                extra_args=extra_args
            )

            # 删除临时文件
            temp_md.unlink()

        except ImportError:
            # 方法2: 使用python-docx手动创建
            print("pypandoc未安装，使用python-docx...")
            self._convert_with_python_docx(markdown_content, output_path)
        except Exception as e:
            # 如果pypandoc失败，也尝试python-docx
            print(f"pypandoc转换失败: {e}，尝试使用python-docx...")
            self._convert_with_python_docx(markdown_content, output_path)

    def _convert_with_python_docx(self, markdown_content, output_path):
        """使用python-docx将Markdown转换为Word（简化版）"""
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        import re

        doc = Document()

        print("\n=== 开始转换Markdown到Word ===")

        # 分行处理markdown内容
        lines = markdown_content.split('\n')
        image_count = 0
        image_not_found = 0

        for i, line in enumerate(lines):
            line = line.rstrip()

            # 处理标题
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('#### '):
                doc.add_heading(line[5:], level=4)
            # 处理图片
            elif '![' in line and '](' in line:
                # 提取所有图片引用
                matches = re.findall(r'!\[([^\]]*)\]\(([^\)]+)\)', line)

                for alt_text, img_path in matches:
                    print(f"\n第{i+1}行发现图片:")
                    print(f"  Alt文本: {alt_text}")
                    print(f"  原始路径: {img_path}")

                    # 如果是相对路径，转换为绝对路径
                    original_img_path = img_path
                    if not Path(img_path).is_absolute():
                        img_path = output_path.parent / img_path

                    print(f"  解析后路径: {img_path}")
                    print(f"  文件是否存在: {Path(img_path).exists()}")

                    if Path(img_path).exists():
                        try:
                            # 获取图片实际大小并适当缩放
                            from PIL import Image as PILImage
                            img = PILImage.open(img_path)
                            width, _ = img.size

                            # 设置最大宽度为6英寸
                            max_width = 6.0
                            doc_width = min(max_width, width / 96)  # 96 DPI

                            doc.add_picture(str(img_path), width=Inches(doc_width))
                            image_count += 1
                            print(f"  ✓ 图片已添加到Word (宽度: {doc_width:.2f}英寸)")
                        except Exception as e:
                            print(f"  ✗ 添加图片失败: {e}")
                            doc.add_paragraph(f"[无法加载图片: {original_img_path}]")
                            import traceback
                            traceback.print_exc()
                    else:
                        print(f"  ✗ 图片文件不存在!")
                        doc.add_paragraph(f"[图片未找到: {original_img_path}]")
                        image_not_found += 1

                # 处理图片后面可能还有的文本
                text_after_img = re.sub(r'!\[([^\]]*)\]\(([^\)]+)\)', '', line).strip()
                if text_after_img:
                    doc.add_paragraph(text_after_img)

            # 处理空行
            elif not line.strip():
                doc.add_paragraph()
            # 普通段落
            else:
                # 简单处理粗体和斜体
                line = re.sub(r'\*\*(.*?)\*\*', r'\1', line)  # 移除粗体标记
                line = re.sub(r'\*(.*?)\*', r'\1', line)  # 移除斜体标记
                doc.add_paragraph(line)

        print(f"\n=== 转换完成 ===")
        print(f"成功添加图片: {image_count}")
        print(f"未找到图片: {image_not_found}")

        doc.save(str(output_path))

    def __del__(self):
        """清理资源"""
        if self.pdf_document:
            self.pdf_document.close()


def main():
    root = tk.Tk()
    app = PDFtoMarkdownConverter(root)
    root.mainloop()


if __name__ == "__main__":
    main()